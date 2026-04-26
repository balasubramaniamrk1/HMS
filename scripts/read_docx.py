import docx

doc = docx.Document("Visiting Doctors Educational with Experience details.docx")
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

for table in doc.tables:
    for row in table.rows:
        print([cell.text.strip() for cell in row.cells])
